import os
import sys
import json
import re
import logging
import google.generativeai as genai
from docx import Document
from PyPDF2 import PdfReader
import markdown

# --- Configuration ---
# CURRENT_VERSION = "1.4 - JSON Output & Final Polish"
# GOAL_VERSION = "2.0 - Perfectly Formatted DOCX Output"

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)

# Configure the Gemini API key
GEMINI_API_KEY = None
ENV_PATH = "/Research_CRM/.env"
if os.path.exists(ENV_PATH):
    with open(ENV_PATH, "r") as f:
        for line in f:
            if line.strip().startswith("GEMINI_API_KEY="):
                GEMINI_API_KEY = (
                    line.strip().split("=", 1)[1].strip().strip('"').strip("'")
                )
                break

if not GEMINI_API_KEY:
    logging.error(f"GEMINI_API_KEY not found in {ENV_PATH}")
    sys.exit(1)

genai.configure(api_key=GEMINI_API_KEY)


# --- File Reading Functions ---
def read_docx_text(path):
    """Reads and returns all text from a .docx file."""
    logging.info(f"Reading DOCX file: {path}")
    try:
        doc = Document(path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        logging.error(f"Error reading docx {path}: {e}")
        return ""


def read_pdf_text(path):
    """Reads and returns all text from a .pdf file."""
    logging.info(f"Reading PDF file: {path}")
    try:
        reader = PdfReader(path)
        full_text = []
        for page in reader.pages:
            full_text.append(page.extract_text())
        return "\n".join(full_text)
    except Exception as e:
        logging.error(f"Error reading pdf {path}: {e}")
        return ""


def get_project_summary(project_name):
    """Reads the Summary.md for a given project."""
    summary_path = os.path.join("projects", project_name, "Summary.md")
    logging.info(f"Reading project summary: {summary_path}")
    if not os.path.exists(summary_path):
        raise FileNotFoundError(
            f"Could not find Summary.md for project: {project_name}"
        )
    with open(summary_path, "r") as f:
        return f.read()


# --- AI Content Generation ---
def generate_section_content(section_title, section_body, project_summary, examples):
    """Uses Gemini to generate the content for a single DSP section."""
    logging.info(f"Preparing prompt for section: '{section_title}'")
    model = genai.GenerativeModel("gemini-2.5-pro")

    relevant_examples = []
    for example in examples:
        if section_title in example:
            start_index = example.find(section_title)
            next_section_match = re.search(
                r"\\n\\n[A-Z][a-z]+", example[start_index + len(section_title) :]
            )
            if next_section_match:
                end_index = (
                    start_index + len(section_title) + next_section_match.start()
                )
                relevant_examples.append(example[start_index:end_index])
            else:
                relevant_examples.append(example[start_index:])
    logging.info(f"Found {len(relevant_examples)} relevant examples for this section.")

    example_texts = "\n--- EXAMPLE ---\n".join(relevant_examples)

    prompt = f"""
You are an expert in writing Data Security Plans for university research. Your task is to complete the following section of a DSP, replacing all placeholders like [PLACEHOLDER] with specific, relevant information.

**Template Section to Complete:**
Title: {section_title}
Body: {section_body}

**Project-Specific Information (from its Summary.md):**
{project_summary}

**Examples of similar, completed sections from other DSPs:**
{example_texts}

**Instructions:**
Based on all the information above, write a new, complete version of the section. It must be comprehensive, professional, and directly relevant to the project's details. **Do not use any placeholder text.** 
Return your response as a single JSON object with one key: "section_content". The value should be the fully completed text for the new section, formatted in Markdown.
"""

    try:
        logging.info(f"Sending request to Gemini API for section: '{section_title}'")
        response = model.generate_content(prompt)

        # Clean and parse the JSON response
        clean_json = response.text.strip().replace("```json", "").replace("```", "")
        parsed_json = json.loads(clean_json)
        content = parsed_json.get(
            "section_content",
            "[[ERROR: 'section_content' key not found in JSON response.]]",
        )

        logging.info(
            f"Successfully received and parsed response for section: '{section_title}'"
        )
        return prompt, content
    except Exception as e:
        logging.error(
            f"Error generating or parsing content for section '{section_title}': {e}"
        )
        return (
            prompt,
            "[[ERROR: Could not generate or parse content for this section. Please review manually.]]",
        )


# --- File Saving Functions ---
def save_generation_log(output_dir, project_name, generation_log):
    """Saves the full generation log as a JSON file."""
    json_path = os.path.join(output_dir, f"{project_name}_DSP_Generation_Log.json")
    logging.info(f"Saving generation log to: {json_path}")
    with open(json_path, "w") as f:
        json.dump(generation_log, f, indent=2)
    logging.info("Generation log saved successfully.")


def save_as_md(output_dir, project_name, generation_log):
    """Saves the generated content as a Markdown file."""
    md_path = os.path.join(output_dir, f"{project_name}_DSP_Generated.md")
    logging.info(f"Saving Markdown output to: {md_path}")
    with open(md_path, "w") as f:
        for item in generation_log:
            f.write(f"# {item['title']}\n\n")
            f.write(f"{item['response']}\n\n")
    logging.info("Markdown file saved successfully.")


def save_as_html(output_dir, project_name, generation_log):
    """Saves the generated content as a modern, interactive HTML file."""
    html_path = os.path.join(output_dir, f"{project_name}_DSP_Generated.html")
    logging.info(f"Saving modern HTML output to: {html_path}")

    # Convert the log to a JSON string for embedding in the HTML
    json_log_str = json.dumps(generation_log, indent=2)

    sections_html = "".join(
        [
            f"""
            <section class="bg-white shadow rounded-lg p-6">
                <h2 class="text-2xl font-semibold text-gray-700 border-b pb-2 mb-4">{item["title"]}</h2>
                <div class="prose max-w-none">
                    {markdown.markdown(item["response"])}
                </div>
            </section>
            """
            for item in generation_log
        ]
    )

    html_template = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Data Security Plan: {project_name}</title>
    <script src="https://cdn.tailwindcss.com"></script>
</head>
<body class="bg-gray-100 font-sans">
    <div class="container mx-auto p-8">
        <header class="bg-white shadow rounded-lg p-6 mb-8">
            <h1 class="text-4xl font-bold text-gray-800">Data Security Plan</h1>
            <p class="text-xl text-gray-600 mt-2">{project_name}</p>
        </header>

        <div class="mb-8">
            <button id="toggle-json-btn" class="bg-blue-500 hover:bg-blue-700 text-white font-bold py-2 px-4 rounded">
                Show/Hide Raw Generation Log
            </button>
            <div id="json-container" class="hidden mt-4 bg-gray-800 text-white p-4 rounded-lg overflow-x-auto">
                <pre><code>{json_log_str}</code></pre>
            </div>
        </div>

        <main class="space-y-8">
            {sections_html}
        </main>
    </div>

    <script>
        document.getElementById('toggle-json-btn').addEventListener('click', function() {{
            var container = document.getElementById('json-container');
            if (container.classList.contains('hidden')) {{
                container.classList.remove('hidden');
            }} else {{
                container.classList.add('hidden');
            }}
        }});
    </script>
</body>
</html>
"""
    with open(html_path, "w") as f:
        f.write(html_template)
    logging.info("Modern HTML file saved successfully.")


# --- Main Execution ---
def main(project_name):
    """Main function to generate the DSP."""
    logging.info(f"--- Starting DSP Generation for project: {project_name} ---")

    # 1. Information Gathering
    logging.info("--- Phase 1: Information Gathering ---")
    try:
        project_summary = get_project_summary(project_name)
    except FileNotFoundError as e:
        logging.error(e)
        return

    try:
        with open("scripts/dsp_template_structure.json", "r") as f:
            structured_template = json.load(f)
        logging.info("Successfully loaded structured template.")
    except FileNotFoundError:
        logging.error(
            "Could not find 'scripts/dsp_template_structure.json'. Please run the structuring script first."
        )
        return

    example_dsp_dir = os.path.join("example_dsps")
    example_texts = []
    for filename in os.listdir(example_dsp_dir):
        if filename.endswith(".docx") and "Template" not in filename:
            example_texts.append(
                read_docx_text(os.path.join(example_dsp_dir, filename))
            )
        elif filename.endswith(".pdf"):
            example_texts.append(read_pdf_text(os.path.join(example_dsp_dir, filename)))
    logging.info(f"Successfully gathered {len(example_texts)} example DSPs.")

    # 2. Iterative Generation
    logging.info("--- Phase 2: Iterative Content Generation (via Gemini) ---")
    generation_log = []

    for i, section in enumerate(structured_template):
        title = section["title"]
        body = section["body"]

        logging.info(f"Processing section {i + 1}/{len(structured_template)}: {title}")

        prompt, response = generate_section_content(
            title, body, project_summary, example_texts
        )
        generation_log.append({"title": title, "prompt": prompt, "response": response})

    # 3. Save the New Documents
    logging.info("--- Phase 3: Saving Output Files ---")
    output_dir = os.path.join("projects", project_name)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        logging.info(f"Created output directory: {output_dir}")

    save_generation_log(output_dir, project_name, generation_log)
    save_as_md(output_dir, project_name, generation_log)
    save_as_html(output_dir, project_name, generation_log)

    logging.info("--- DSP Generation Complete ---")
    print("\nFull DSP draft and logs saved in MD, HTML, and JSON formats.")
    print("Please review the documents for completeness and accuracy.")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python3 dsp_generator_final.py <ProjectFolderName>")
    else:
        main(sys.argv[1])
