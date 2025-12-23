import logging
from docx import Document # type: ignore
from PyPDF2 import PdfReader # type: ignore

logger = logging.getLogger(__name__)

def read_docx_text(path: str) -> str:
    """Reads and returns all text from a .docx file."""
    logger.debug(f"Reading DOCX file: {path}")
    try:
        doc = Document(path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error reading docx {path}: {e}")
        return ""

def read_pdf_text(path: str) -> str:
    """Reads and returns all text from a .pdf file."""
    logger.debug(f"Reading PDF file: {path}")
    try:
        reader = PdfReader(path)
        full_text = []
        for page in reader.pages:
            full_text.append(page.extract_text())
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error reading pdf {path}: {e}")
        return ""

def read_file_content(path: str) -> str:
    """Dispatches to the correct reader based on extension."""
    if path.endswith(".docx"):
        return read_docx_text(path)
    elif path.endswith(".pdf"):
        return read_pdf_text(path)
    elif path.endswith(".md") or path.endswith(".txt"):
         with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        logger.warning(f"Unsupported file type: {path}")
        return ""
